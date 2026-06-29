import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()
    
    # Page settings - Standard Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Styling helper functions
    def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def add_styled_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        for run in heading.runs:
            set_font(run, name="Calibri", size=18 - (level * 2), bold=True, color=RGBColor(13, 95, 132))
        return heading

    def add_body_paragraph(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_pref = p.add_run(bold_prefix)
            set_font(run_pref, size=11, bold=True)
        run_text = p.add_run(text)
        set_font(run_text, size=11)
        return p

    def add_bullet_point(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_pref = p.add_run(bold_prefix)
            set_font(run_pref, size=11, bold=True)
        run_text = p.add_run(text)
        set_font(run_text, size=11)
        return p

    # ------------------ COVER PAGE ------------------
    for _ in range(5):
        doc.add_paragraph() # Spacing
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("LAPORAN PENGUJIAN PERANGKAT LUNAK\nAPLIKASI HOMI (LAYANAN WARGA)")
    set_font(run_title, size=24, bold=True, color=RGBColor(13, 95, 132))
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("Pengujian Sistematis: Playwright E2E, Postman/Newman API, dan k6 Performance")
    set_font(run_sub, size=12, italic=True)
    
    for _ in range(8):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Disusun untuk Mata Kuliah: Kapita Selekta\nProgram Studi Teknologi Informasi\n\nTahun 2026")
    set_font(run_meta, size=11, bold=True)
    
    doc.add_page_break()

    # ------------------ SECTION A ------------------
    add_styled_heading("A. Pendahuluan/Tujuan Pengujian", 1)
    add_body_paragraph(
        "Aplikasi HOMI merupakan platform sistem informasi layanan warga berbasis multi-tenancy. "
        "Tujuan utama dari pengujian sistematis ini adalah menjamin fungsionalitas dan kualitas aplikasi sebelum rilis. "
        "Fokus pengujian meliputi:"
    )
    add_bullet_point("Memastikan modul-modul penting web admin HOMI dapat diakses dan berjalan dengan stabil.")
    add_bullet_point("Memverifikasi isolasi data multi-tenant agar data antar perumahan tidak saling bocor.")
    add_bullet_point("Menguji performa backend API di bawah beban simultan ringan (load testing).")
    add_bullet_point("Menemukan celah konfigurasi dan bug pada lingkungan pengembangan lokal.")

    # ------------------ SECTION B ------------------
    add_styled_heading("B. Tools yang Digunakan", 1)
    add_body_paragraph(
        "Pengujian ini mengintegrasikan beberapa perkakas pengujian modern untuk mencakup berbagai aspek kualitas:"
    )
    add_bullet_point("E2E testing antarmuka web admin, visual check, dan otomasi tangkapan layar (screenshot).", "Playwright: ")
    add_bullet_point("Pengujian fungsional API, kepatuhan status code, dan validasi tenancy header.", "Postman / Newman: ")
    add_bullet_point("Load testing API backend ringan secara lokal untuk memantau responsivitas server.", "k6: ")
    add_bullet_point("Database server lokal tempat data tenant terisolasi.", "MySQL 8.4 (Laragon): ")

    # ------------------ SECTION C ------------------
    add_styled_heading("C. Skenario Pengujian", 1)
    add_bullet_point("Autentikasi admin perumahan (Hawaii Garden) dengan skenario sukses dan gagal.", "Skenario 1 - Login & Auth: ")
    add_bullet_point("Memuat Dashboard, Data Warga, Tagihan, Pembayaran, Pengumuman, Pengaduan, dan ranking Prioritas SAW.", "Skenario 2 - Navigasi Modul: ")
    add_bullet_point("Memastikan visualisasi antarmuka dan data yang dimuat terbatas hanya pada scope perumahan Hawaii Garden.", "Skenario 3 - Proteksi Multi-Tenant: ")
    add_bullet_point("Mengirim request API dengan header X-Tenant-Code yang salah, kosong, maupun valid untuk menguji middleware ResolveTenant.", "Skenario 4 - API Tenancy: ")
    add_bullet_point("Simulasi 5 Virtual Users secara konkuren mengakses API ping dan list tenants selama 30 detik.", "Skenario 5 - Load Test k6: ")

    # ------------------ SECTION D ------------------
    add_styled_heading("D. Langkah-langkah Pengujian", 1)
    add_bullet_point("Pastikan database MySQL Laragon berjalan dan dikonfigurasi di berkas `.env`.")
    add_bullet_point("Jalankan migrasi database pusat dan seluruh tenant menggunakan `php artisan migrate` dan `php artisan homi:tenants-migrate`.")
    add_bullet_point("Inisialisasi tenant Hawaii Garden dengan `php artisan tenant:initialize hawaii-garden` dan sinkronkan admin dengan `php artisan homi:sync-admins`.")
    add_bullet_point("Jalankan server pengembangan Laravel (`php artisan serve`) dan aset front-end (`npm run dev`).")
    add_bullet_point("Jalankan test suite Playwright untuk web E2E (`npx playwright test`).")
    add_bullet_point("Jalankan collection API via Newman (`npx newman run ...`).")
    add_bullet_point("Jalankan simulasi beban API via k6 (`k6 run ...`).")

    # ------------------ SECTION E ------------------
    add_styled_heading("E. Hasil Pengujian", 1)
    
    # E.1 Tabel Test Case Pengujian
    add_styled_heading("E.1 Tabel Test Case Pengujian", 2)
    add_body_paragraph(
        "Seluruh skenario pengujian fungsionalitas visual web admin berhasil diselesaikan. "
        "Tabel berikut merangkum hasil uji kasus (test cases) lengkap. Detail screenshot visual diletakkan pada Lampiran A."
    )

    # Table of Test Cases (Font Size 8 pt to avoid overflow)
    headers = ["ID", "Fitur", "Tools", "Skenario", "Expected Result", "Actual Result", "Status"]
    test_cases_data = [
        ["TC001", "Auth", "Playwright", "Login sukses", "Masuk ke Dashboard Hawaii Garden", "Redirect sukses", "PASSED"],
        ["TC002", "Auth", "Playwright", "Login gagal", "Ditolak & muncul error di form", "Ditolak dengan pesan error", "PASSED"],
        ["TC003", "Dash", "Playwright", "Dashboard", "Tampil ringkasan data tenant", "Widget ringkasan dimuat", "PASSED"],
        ["TC004", "Warga", "Playwright", "Data Warga", "Tampil list warga Hawaii Garden", "List tabel termuat", "PASSED"],
        ["TC005", "Tagihan", "Playwright", "Tagihan Iuran", "Tampil riwayat invoice perumahan", "Invoice termuat", "PASSED"],
        ["TC006", "Bayar", "Playwright", "Pembayaran", "Tampil riwayat pembayaran warga", "Tabel transaksi dimuat", "PASSED"],
        ["TC007", "Pengumuman", "Playwright", "Pengumuman", "Tampil daftar informasi pengumuman", "Modul ter-render", "PASSED"],
        ["TC008", "Aduan", "Playwright", "Pengaduan", "Tampil daftar keluhan warga", "Aspirasi dimuat", "PASSED"],
        ["TC009", "SAW", "Playwright", "Prioritas SAW", "Tampil perangkingan tunggakan", "Skor SAW alternatif dimuat", "PASSED"],
        ["TC010", "Tenancy", "Playwright", "Multi-Tenant", "Hanya data Hawaii Garden yang tampil", "Data scope Hawaii Garden saja", "PASSED"]
    ]

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Format Headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_font(run, size=8.5, bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D5F84"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Populate Rows
    for row_idx, row_data in enumerate(test_cases_data):
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            p = row_cells[col_idx].paragraphs[0]
            if col_idx in [0, 2, 6]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, size=8)
                if text == "PASSED":
                    run.bold = True
                    run.font.color.rgb = RGBColor(46, 125, 50)
            
            if row_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F8FA"/>')
                row_cells[col_idx]._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph() # spacing

    # E.2 Bukti Eksekusi Tools
    add_styled_heading("E.2 Bukti Eksekusi Tools", 2)
    add_body_paragraph(
        "Seluruh bukti eksekusi nyata berupa log keluaran terminal telah disimpan dan di-render secara visual "
        "pada Lampiran B. File log teks lengkap dapat diakses pada direktori proyek:"
    )
    add_bullet_point("Logs Playwright: docs/testing/logs/playwright-result.txt")
    add_bullet_point("Logs Postman/Newman: docs/testing/logs/api-test-result.txt")
    add_bullet_point("Logs k6: docs/testing/logs/k6-result.txt")

    # E.3 Ringkasan Hasil Pengujian
    add_styled_heading("E.3 Ringkasan Hasil Pengujian", 2)
    add_body_paragraph(
        "Ringkasan performa dan kestabilan aplikasi setelah perbaikan konfigurasi local development:"
    )
    add_bullet_point("Playwright E2E: 100% Passed. Alur login admin dan semua modul utama berhasil terbuka tanpa ada error.", "E2E Web Admin: ")
    add_bullet_point("Postman/Newman: 13 asersi passed, 0 failed. Proteksi middleware tenant berhasil mengidentifikasi dan memblokir token tenant silang maupun request tanpa header.", "API Backend: ")
    add_bullet_point("k6 Load Testing: 0% request failed dengan rata-rata respon di bawah 300ms. Kestabilan konkurensi lokal meningkat drastis setelah Redis dinonaktifkan.", "Load Performance: ")

    # ------------------ SECTION F ------------------
    add_styled_heading("F. Temuan Bug/Issue", 1)
    add_body_paragraph(
        "1. Ketergantungan Redis secara Default (Configuration Bug):", bold_prefix=""
    )
    add_body_paragraph(
        "Berkas `.env` default mewajibkan Redis (`CACHE_STORE=redis`) sehingga memicu error 500 jika Redis server tidak aktif secara lokal. "
        "Solusi dilakukan dengan mengubah parameter tersebut ke `file` untuk kemudahan local setup."
    )
    add_body_paragraph(
        "2. Keterbatasan Concurrency PHP CLI Server (Performance Bottleneck):", bold_prefix=""
    )
    add_body_paragraph(
        "Pengujian k6 dengan 5 VUs konkuren menghasilkan kerentanan bottleneck pada server development default `php artisan serve` yang bersifat single-threaded apabila dibebani request paralel padat."
    )
    add_body_paragraph(
        "3. Ketiadaan Endpoint API untuk Dashboard Summary dan SAW (Missing API Routes):", bold_prefix=""
    )
    add_body_paragraph(
        "Fitur dashboard summary dan perangkingan SAW belum diposisikan sebagai endpoint API JSON di `api.php`, melainkan baru terpasang di rute web biasa."
    )

    # ------------------ SECTION G ------------------
    add_styled_heading("G. Analisis dan Rekomendasi Perbaikan", 1)
    add_body_paragraph(
        "1. Pengaturan Cache Default:", bold_prefix=""
    )
    add_body_paragraph("Mengubah nilai default `CACHE_STORE` di `.env.example` ke `file` agar developer baru dapat langsung menjalankan aplikasi tanpa Redis.")
    add_body_paragraph(
        "2. Peningkatan Web Server lokal untuk Performance Test:", bold_prefix=""
    )
    add_body_paragraph("Hasil benchmark k6 menunjukkan performa server local development. Disarankan menggunakan Nginx + PHP-FPM di Laragon atau Laravel Octane untuk performa konkuren.")
    add_body_paragraph(
        "3. Pemisahan Endpoint Layanan Mobile:", bold_prefix=""
    )
    add_body_paragraph("Membuat API khusus berformat JSON untuk fitur dashboard summary dan SAW agar aplikasi Android dapat menarik data tersebut dengan rapi.")

    # ------------------ SECTION H ------------------
    add_styled_heading("H. Kesimpulan", 1)
    add_body_paragraph(
        "Berdasarkan skenario pengujian yang dilakukan, fitur utama web admin HOMI pada scope tenant Hawaii Garden berhasil dijalankan dengan baik. "
        "Data yang tampil berada dalam scope tenant Hawaii Garden dan tidak ditemukan indikasi data tenant lain muncul selama pengujian berlangsung. "
        "Namun, pengujian ini masih terbatas pada skenario internal dan belum dapat dianggap sebagai audit keamanan menyeluruh."
    )

    doc.add_page_break()

    # ------------------ LAMPIRAN A ------------------
    add_styled_heading("Lampiran A: Screenshot E2E TC001–TC010", 1)
    
    screenshots = [
        ("TC001_login_admin_success.png", "Bukti pengujian login admin berhasil"),
        ("TC002_login_failed_wrong_password.png", "Bukti pengujian login gagal jika password salah"),
        ("TC003_dashboard_loaded.png", "Bukti pengujian memuat halaman Dashboard"),
        ("TC004_resident_page_loaded.png", "Bukti pengujian memuat halaman Data Warga"),
        ("TC005_invoice_page_loaded.png", "Bukti pengujian memuat halaman Tagihan Iuran"),
        ("TC006_payment_page_loaded.png", "Bukti pengujian memuat halaman Pembayaran"),
        ("TC007_announcement_page_loaded.png", "Bukti pengujian memuat halaman Pengumuman"),
        ("TC008_complaint_page_loaded.png", "Bukti pengujian memuat halaman Pengaduan"),
        ("TC009_saw_priority_page_loaded.png", "Bukti pengujian memuat halaman prioritas tunggakan"),
        ("TC010_multi_tenant_hawaii_garden.png", "Bukti pengujian keamanan data multi-tenant Hawaii Garden")
    ]
    
    e2e_dir = "docs/testing/screenshots/e2e"
    for filename, caption in screenshots:
        filepath = os.path.join(e2e_dir, filename)
        if os.path.exists(filepath):
            doc.add_paragraph()
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(filepath, width=Inches(4.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(f"Gambar. {caption}")
            set_font(run_cap, size=9.5, italic=True)
            
    doc.add_page_break()

    # ------------------ LAMPIRAN B ------------------
    add_styled_heading("Lampiran B: Bukti Eksekusi Tools", 1)
    
    tools_screenshots = [
        ("playwright_result.png", "Bukti eksekusi Playwright E2E Test Suite (2 Passed)"),
        ("api_test_result.png", "Bukti eksekusi Newman API Testing (13 Assertions Passed)"),
        ("k6_result.png", "Bukti eksekusi k6 Performance Load Test (0% Failures)")
    ]
    
    tools_dir = "docs/testing/screenshots/tools"
    for filename, caption in tools_screenshots:
        filepath = os.path.join(tools_dir, filename)
        if os.path.exists(filepath):
            doc.add_paragraph()
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(filepath, width=Inches(4.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(f"Gambar. {caption}")
            set_font(run_cap, size=9.5, italic=True)

    # Save Document
    out_dir = "docs/testing"
    os.makedirs(out_dir, exist_ok=True)
    filepath = os.path.join(out_dir, "Laporan_Pengujian_Perangkat_Lunak_HOMI_FINAL.docx")
    doc.save(filepath)
    print(f"Document saved to {filepath}")

if __name__ == "__main__":
    create_document()
